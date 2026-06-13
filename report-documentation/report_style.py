"""보고서 docx 디자인 토큰과 빌더 헬퍼.

확정 디자인: A 기반 모던 학술형 · Pretendard · 네이비 #1F3A5F.
스펙: docs-web/docs/superpowers/specs/2026-06-13-report-docx-design.md
"""
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# ---- 색상 토큰 ----
NAVY = "1F3A5F"   # 주색
INK = "1A2230"    # 제목
BODY = "2A2F3A"   # 본문
MUTE = "8A93A1"   # 캡션·페이지번호
ZEBRA = "F6F8FA"  # 표 짝수행
HAIR = "E7EAEE"   # 칸선·구분선
WHITE = "FFFFFF"

# ---- Pretendard 가족명(설치 굵기) ----
F_REG = "Pretendard"
F_MED = "Pretendard Medium"
F_SEMI = "Pretendard SemiBold"
F_BOLD = "Pretendard"          # bold 플래그와 함께 사용
F_XBOLD = "Pretendard ExtraBold"

CONTENT_W = Mm(170)  # A4(210) - 좌우 여백(20*2)


# ============================ 저수준 헬퍼 ============================
def set_run_font(run, family, size_pt, color_hex=None, bold=False, spacing_pt=None):
    """run에 한/영 동일 글꼴·크기·색·자간을 설정한다."""
    run.font.name = family
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), family)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if spacing_pt is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing_pt * 20)))  # 1pt = 20 (twentieths of a pt)
        rpr.append(sp)


def _border_el(edge, size_pt, color_hex, space_pt=4):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(int(size_pt * 8)))   # sz: 1/8 pt
    el.set(qn("w:space"), str(int(space_pt)))
    el.set(qn("w:color"), color_hex)
    return el


def set_par_border(paragraph, edge, size_pt, color_hex, space_pt=4):
    """문단에 한 변 테두리(예: H2 좌측 바, 하단 규칙선)를 추가한다."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    pbdr.append(_border_el(edge, size_pt, color_hex, space_pt))


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_borders(cell, **edges):
    """edges 예: top=(0.75, HAIR). 지정한 변만 그린다."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcb = tcPr.find(qn("w:tcBorders"))
    if tcb is None:
        tcb = OxmlElement("w:tcBorders")
        tcPr.append(tcb)
    for edge, (size_pt, color) in edges.items():
        tcb.append(_border_el(edge, size_pt, color, space_pt=0))


def clear_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def set_row_height(row, pt, exact=True):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(pt * 20)))
    h.set(qn("w:hRule"), "exact" if exact else "atLeast")
    trPr.append(h)


def add_page_number(paragraph, family=F_REG, size_pt=9, color=MUTE):
    """문단에 PAGE 필드를 넣는다(현재 페이지 번호)."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), family)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    rpr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rpr.append(col)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def restart_page_numbering(section, start=1):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:start"), str(start))


def setup_page(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)


def init_document(doc):
    """기본(Normal) 글꼴을 Pretendard로 고정."""
    normal = doc.styles["Normal"]
    normal.font.name = F_REG
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), F_REG)
    setup_page(doc.sections[0])


# ============================ 컴포넌트 빌더 ============================
def _blank(doc, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def accent_bar(doc, width_mm=15, height_pt=3, color=NAVY, align=WD_TABLE_ALIGNMENT.CENTER):
    """중앙 정렬 네이비 강조 바(얇은 사각형)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = align
    clear_table_borders(t)
    t.columns[0].width = Mm(width_mm)
    cell = t.cell(0, 0)
    cell.width = Mm(width_mm)
    shade_cell(cell, color)
    set_row_height(t.rows[0], height_pt, exact=True)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    set_run_font(run, F_REG, 1, color)
    return t


def heading1(doc, kicker, title):
    """장 제목: 킥커(제N장) + 큰 제목 + 하단 규칙선. 새 페이지에서 호출 가정."""
    pk = doc.add_paragraph()
    pk.paragraph_format.space_after = Pt(2)
    pk.paragraph_format.space_before = Pt(6)
    set_run_font(pk.add_run(kicker), F_XBOLD, 12, NAVY)
    pt = doc.add_paragraph()
    pt.paragraph_format.space_after = Pt(10)
    set_run_font(pt.add_run(title), F_XBOLD, 23, INK, spacing_pt=-0.2)
    set_par_border(pt, "bottom", 0.75, HAIR, space_pt=10)
    return pt


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(9)
    set_par_border(p, "left", 3, NAVY, space_pt=8)
    set_run_font(p.add_run(text), F_BOLD, 15, NAVY, bold=True)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(text), F_SEMI, 12.5, INK)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.75
    set_run_font(p.add_run(text), F_REG, 10.5, BODY)
    return p


def caption(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(label + " "), F_BOLD, 10, NAVY, bold=True)
    set_run_font(p.add_run(text), F_REG, 10, MUTE)
    return p


def data_table(doc, header, rows):
    """네이비 채움 헤더 + 줄무늬 표."""
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    clear_table_borders(t)
    # 헤더
    for i, htext in enumerate(header):
        c = t.cell(0, i)
        shade_cell(c, NAVY)
        set_cell_borders(c, bottom=(0.75, NAVY))
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(htext), F_BOLD, 10, WHITE, bold=True)
    # 본문 행
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        even = (ri % 2 == 1)
        for ci, val in enumerate(row):
            c = cells[ci]
            if even:
                shade_cell(c, ZEBRA)
            set_cell_borders(c, bottom=(0.75, HAIR))
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(val), F_REG, 10, BODY)
    return t


def toc_entry(doc, num, title, page):
    """목차 한 줄: 번호 + 제목 + 점선 리더 + 페이지번호(우측)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3.5)
    pf.space_after = Pt(3.5)
    pf.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if num:
        set_run_font(p.add_run(num + "  "), F_XBOLD, 11, NAVY)
    set_run_font(p.add_run(title), F_SEMI, 11, INK)
    set_run_font(p.add_run("\t" + str(page)), F_REG, 11, MUTE)
    return p
