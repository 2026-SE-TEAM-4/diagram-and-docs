"""표지 + 목차 + 디자인 데모가 들어간 보고서 docx 초안을 생성한다.

사용:  python build_report.py
산출:  서버예약할당관리시스템_보고서_초안.docx
스펙:  docs-web/docs/superpowers/specs/2026-06-13-report-docx-design.md
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Mm, Pt

import report_style as S

OUT = Path(__file__).with_name("서버예약할당관리시스템_보고서_초안.docx")

# 자리표시(확정 시 교체)
UNIV = "한국공학대학교 · 소프트웨어공학"
KICKER = "2026학년도 1학기 팀 프로젝트"
TITLE_L1 = "서버 예약/할당"
TITLE_L2 = "관리 시스템"
SUBTITLE = ["예약 기반 자원 할당과", "AIOps 예측 운영"]
COVER_INFO = [
    ("팀", "소프트웨어공학 4조"),
    ("팀원", "○○○(팀장) · ○○○ · ○○○"),
    ("제출일", "2026. 06."),
]

# 목차 (번호, 제목, 자리표시 페이지)
TOC = [
    ("", "요약", 1),
    ("Ⅰ", "개발 목표·배경 및 필요성", 2),
    ("Ⅱ", "유사 시스템 사례 분석", 5),
    ("Ⅲ", "개발 프로세스·프로젝트 계획", 8),
    ("Ⅳ", "요구사항 분석", 11),
    ("Ⅴ", "아키텍처 분석·전략설계", 16),
    ("Ⅵ", "상세설계 — 정적 모델", 20),
    ("Ⅶ", "상세설계 — 자료구조·알고리즘", 24),
    ("Ⅷ", "SOLID 원칙 적용", 28),
    ("Ⅸ", "디자인 패턴 적용", 31),
    ("Ⅹ", "테스트", 34),
    ("Ⅺ", "구현 결과", 38),
    ("Ⅻ", "리스크 및 느낀점", 41),
    ("", "부록", 44),
]


def build_cover(doc):
    # 상단: 학교·과목 + 하단 가는 규칙선
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    S.set_run_font(p.add_run(UNIV), S.F_BOLD, 11.5, S.NAVY, bold=True)
    S.set_par_border(p, "bottom", 0.75, S.HAIR, space_pt=10)

    # 중앙: 킥커 → 제목 → 강조 바 → 부제
    pk = doc.add_paragraph()
    pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pk.paragraph_format.space_before = Pt(118)
    S.set_run_font(pk.add_run(KICKER), S.F_MED, 11, S.MUTE)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(12)
    pt.paragraph_format.line_spacing = 1.22
    r1 = pt.add_run(TITLE_L1)
    S.set_run_font(r1, S.F_XBOLD, 28, S.INK, spacing_pt=-0.4)
    r1.add_break()
    r2 = pt.add_run(TITLE_L2)
    S.set_run_font(r2, S.F_XBOLD, 28, S.INK, spacing_pt=-0.4)

    S.accent_bar(doc, width_mm=15, height_pt=3)

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(12)
    ps.paragraph_format.line_spacing = 1.5
    rs1 = ps.add_run(SUBTITLE[0])
    S.set_run_font(rs1, S.F_MED, 12, "5B6675")
    rs1.add_break()
    rs2 = ps.add_run(SUBTITLE[1])
    S.set_run_font(rs2, S.F_MED, 12, "5B6675")

    # 하단: 정보표 (페이지 하단으로 밀기 위한 spacer)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(250)

    t = doc.add_table(rows=len(COVER_INFO), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    S.clear_table_borders(t)
    t.columns[0].width = Mm(30)
    t.columns[1].width = Mm(140)
    for ri, (k, v) in enumerate(COVER_INFO):
        kc, vc = t.cell(ri, 0), t.cell(ri, 1)
        kc.width, vc.width = Mm(30), Mm(140)
        top = (1.5, S.NAVY) if ri == 0 else (0.75, S.HAIR)
        S.set_cell_borders(kc, top=top)
        S.set_cell_borders(vc, top=top)
        kp = kc.paragraphs[0]
        kp.paragraph_format.space_before = Pt(5)
        kp.paragraph_format.space_after = Pt(5)
        S.set_run_font(kp.add_run(k), S.F_SEMI, 11, S.MUTE)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vp.paragraph_format.space_before = Pt(5)
        vp.paragraph_format.space_after = Pt(5)
        S.set_run_font(vp.add_run(v), S.F_SEMI, 11, "2A3442")


def build_toc(doc):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(2)
    S.set_run_font(h.add_run("목차"), S.F_XBOLD, 18, S.INK)
    S.set_par_border(h, "bottom", 2, S.NAVY, space_pt=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    for num, title, page in TOC:
        S.toc_entry(doc, num, title, page)


def build_demo(doc):
    """디자인 적용 확인용 — 제Ⅵ장 일부를 실제 스타일로 렌더."""
    S.heading1(doc, "제6장", "상세설계 — 정적 모델")
    S.heading2(doc, "6.1 레이어 구조")
    S.body(doc, "백엔드는 api · core · services · models · jobs 다섯 레이어로 구성된다. "
                "라우터는 얇게 유지하고 비즈니스 로직은 services에 둔다. 각 레이어는 한 방향으로만 "
                "의존하여 변경의 파급을 줄인다.")
    S.heading3(doc, "6.1.1 ORM 엔티티 17개")
    S.body(doc, "도메인은 17개 엔티티로 모델링하며, 예약(Reservation)·서버(Server)·한도(Quota)가 "
                "핵심 축이다. 동시성은 Server.version 낙관적 잠금으로 제어한다.")
    S.data_table(
        doc,
        ["엔티티", "역할", "키"],
        [
            ["Reservation", "예약 단위 — 사용자·서버·기간", "PK · FK"],
            ["Server", "자원과 상태(가용·예약·사용중·점검)", "PK"],
            ["Quota", "사용자별 예약 한도와 사용량", "PK · FK"],
        ],
    )
    S.caption(doc, "표 6-1.", "주요 엔티티")
    S.heading2(doc, "6.2 엔티티 관계 (ERD)")
    S.body(doc, "전체 17개 엔티티의 관계는 ERD(그림 6-1)로 나타낸다. 식별/비식별 관계와 카디널리티를 "
                "함께 표기하며, 외래키 행은 옅은 초록으로 구분한다.")
    S.caption(doc, "그림 6-1.", "전체 ERD (17 엔티티) — 자리표시")


def add_body_header_footer(section):
    """본문 섹션: 좌 시스템명 · 우 장 제목 머리말 + 가운데 페이지번호 꼬리말."""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    hp = section.header.paragraphs[0]
    hp.paragraph_format.tab_stops.add_tab_stop(S.CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
    S.set_run_font(hp.add_run("서버 예약/할당 관리 시스템"), S.F_REG, 9, S.MUTE)
    S.set_run_font(hp.add_run("\t제6장 상세설계 — 정적 모델"), S.F_REG, 9, S.MUTE)
    S.set_par_border(hp, "bottom", 0.75, S.HAIR, space_pt=6)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    S.add_page_number(fp, family=S.F_REG, size_pt=9, color=S.MUTE)


def main():
    doc = Document()
    S.init_document(doc)

    # 섹션0: 표지 (머리말/꼬리말 없음)
    build_cover(doc)

    # 섹션1: 목차 (머리말/꼬리말 없음)
    sec_toc = doc.add_section(WD_SECTION.NEW_PAGE)
    S.setup_page(sec_toc)
    build_toc(doc)

    # 섹션2: 본문 (머리말/꼬리말 + 페이지번호 1부터)
    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    S.setup_page(sec_body)
    S.restart_page_numbering(sec_body, start=1)
    add_body_header_footer(sec_body)
    build_demo(doc)

    doc.save(OUT)
    print(f"saved: {OUT}  ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
